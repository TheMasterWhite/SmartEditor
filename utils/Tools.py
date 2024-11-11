from datetime import datetime
import os
import uuid
import time
import jwt
import logging
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GLOBAL_RSA_PUBLIC_KEY = os.environ.get("GLOBAL_RSA_PUBLIC_KEY")


# 获取格式化时间
def GetTime():
    curTime = datetime.now()
    formatTime = curTime.strftime("%H:%M:%S")
    return str(formatTime)


# 获取格式化日期
def GetDate():
    curDate = datetime.now()
    formatDate = curDate.strftime("%y.%m.%d")
    return str(formatDate)


# 获取文件扩展名
def GetExtension(FileName):
    fileName, extension = os.path.splitext(FileName)
    return extension[1:] if extension.startswith('.') else extension


# 获取文件名
def GetFileName(FileName):
    fileName, extension = os.path.splitext(FileName)
    return fileName


# 获取文件保存时的格式化时间结构
def GetSaveTime():
    curDate = datetime.now().strftime("%y.%m.%d")
    curTime = datetime.now().strftime("%H:%M:%S")
    saveTime = f"{curDate} {curTime}"
    return saveTime


# 验证密码合法性
def ValidPassword(Password):
    if len(Password) < 8 or len(Password) > 16:
        return False
    has_upper = any(char.isupper() for char in Password)
    has_lower = any(char.islower() for char in Password)
    has_digit = any(char.isdigit() for char in Password)
    return has_upper and has_lower and has_digit and Password.isalnum()


# 验证用户名合法性
def ValidUsername(Username):
    if len(Username) < 3 or len(Username) > 10:
        return False
    return Username.isalnum()


# 根据时间生成一个uuid
def GetUUID():
    timestamp = int(time.time())
    UUID = uuid.uuid1()
    return str(UUID)


# 保存到docx中
def SaveDocx(Title, SavePath, Content):
    try:
        doc = Document()
        # 设置中文标题样式
        titleStyle = doc.styles.add_style("DocTitle", 1)
        titleFont = titleStyle.font
        titleFont.name = "SimSun"
        titleFont.size = Pt(22)  # 二号字体
        titleParagraphFormat = titleStyle.paragraph_format
        titleParagraphFormat.space_after = Pt(0)  # 标题后无空行
        doc.add_paragraph(Title, style = "DocTitle")

        # 设置中文正文样式
        bodyStyle = doc.styles["Normal"]  # 正文样式
        bodyFont = bodyStyle.font
        bodyFont.name = "SimSun"
        bodyFont.size = Pt(10.5)  # 五号字体
        doc.add_paragraph(Content, style = "Normal")
        doc.save(SavePath)

    except Exception as e:
        curTime = GetTime()
        logging.error(f"[{curTime}]Module:[SaveDocx]" + str(e))
        raise e
